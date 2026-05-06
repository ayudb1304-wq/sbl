"""
Generate docs/SBL_Feature_List.docx — a structured feature inventory of the
app as it stands today. Re-run any time the feature set evolves.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parents[1] / "docs" / "SBL_Feature_List.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(0x1E, 0x2A, 0x6E)
GREY = RGBColor(0x76, 0x70, 0x6A)
TEXT = RGBColor(0x1F, 0x1B, 0x16)


def add_run(p, text, *, bold=False, italic=False, size=11, color=TEXT):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return r


def add_heading(doc, text, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    add_run(p, text, bold=True, size=sizes.get(level, 12), color=NAVY)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, *, italic=False, color=TEXT):
    p = doc.add_paragraph()
    add_run(p, text, italic=italic, color=color)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, text)
    return p


def add_section_block(doc, title, intro, bullets):
    add_heading(doc, title, level=2)
    if intro:
        add_para(doc, intro, italic=True, color=GREY)
    for b in bullets:
        add_bullet(doc, b)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        add_run(p, h, bold=True, color=NAVY)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            add_run(cell.paragraphs[0], val)
    doc.add_paragraph()


def main():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title, "Sysfore Badminton League", bold=True, size=28, color=NAVY)
    sub = doc.add_paragraph()
    add_run(sub, "Tournament tracking app — feature inventory", italic=True, color=GREY, size=12)
    meta = doc.add_paragraph()
    add_run(meta, "Built on Next.js 16 + Supabase · multi-year (SBL 2026 →)", color=GREY, size=10)

    add_para(doc,
        "This document inventories every shipped feature, grouped by audience and "
        "subsystem. It's the canonical reference when onboarding new contributors, "
        "drafting release notes, or scoping next-iteration work.",
    )

    # ---------- Personas ----------
    add_heading(doc, "Audiences", level=1)
    add_table(doc,
        ["Persona", "Account?", "Capabilities"],
        [
            ["Participant", "No — public read", "Browse the entire tournament: live scores, standings, fixtures, brackets, court schedules, team and player pages. Engage anonymously: follow teams, send cheers, make bracket-challenge picks, climb the leaderboard."],
            ["Scorer", "Email login (whitelisted)", "Enter scores per game, start matches, end games, set winners, declare walkovers, reset matches. All actions logged."],
            ["Admin", "Email login (whitelisted, can also hold scorer role)", "Everything a scorer can do, plus: confirm group qualifiers, generate / re-resolve KO bracket, lock/unlock matches, force-edit completed matches, manage allowed-user whitelist with multi-role, create and activate seasons, view full per-match audit log."],
        ],
    )

    # ---------- Public / participant ----------
    add_heading(doc, "Participant features", level=1)

    add_section_block(doc,
        "Home dashboard",
        "Single screen overview of the whole tournament.",
        [
            "Hero panel: SBL logo, season chip with live status, sporty tagline, match counters (total / completed / live).",
            "Your teams rail (localStorage-backed follow) — pinned teams show next/live match + last result without an account.",
            "Live now: cards for every match currently in progress, status-tinted red.",
            "Up next: the next 6 scheduled matches across all courts.",
            "Recent results: 6 most recently completed matches.",
            "Standings snapshot — auto-rotating carousel cycling every group within each category every 5 seconds; manual prev/next chevrons + dot navigation; tapping pauses auto-rotate.",
        ],
    )

    add_section_block(doc,
        "Browsing",
        "Drill into any slice of the tournament.",
        [
            "Categories (MB / MI / W) — full standings + fixture cards per group.",
            "Group detail — ranked standings with tie-break columns (P, W, L, SD, PD, Pts), team list with players, all fixtures.",
            "Team detail — players, all fixtures, plus a Follow toggle.",
            "Player detail — cross-season tournament history (all teams across all SBL years).",
            "Match detail — score per game, status pill, court / time, walkover reason if applicable, full game-by-game breakdown.",
            "Court page — schedule per court with chip-style court switcher.",
            "Bracket page — visual QF → SF → Final flow per category, with feeder labels (\"MB-A Winner\" etc.) until teams resolve. No horizontal scroll required.",
        ],
    )

    add_section_block(doc,
        "Engagement pack",
        "Anonymous interactivity tied to a localStorage device UUID.",
        [
            "Follow teams (★) — pin teams to a personal Your-Teams rail on home; persists across sessions.",
            "Universal search — Cmd-K (or click) opens a fuzzy-match modal across all teams + players, with keyboard navigation.",
            "Cheers — 👏 Clap / 🔥 Fire buttons on every live or completed match. Each tap is logged anonymously, aggregate counts update in realtime via Supabase channels, with floating emoji animation on tap.",
            "Bracket challenge — pick the winner of every KO match (1 pt QF, 2 pt SF, 4 pt Final). Picks lock when the match starts. Cards turn green on correct picks, red on wrong.",
            "Leaderboard — aggregated correct picks + points across all participants, with display name set via inline NameGate (no signup).",
        ],
    )

    add_section_block(doc,
        "Realtime + always-on score ticker",
        "Every page reflects the live tournament state without manual refresh.",
        [
            "Top marquee — slim navy stripe across the top of every page scrolling live and recent matches; click to jump to a match; hover to pause; sticky LIVE · N badge with pulsing dot.",
            "Realtime score subscriptions on matches and games tables across all key pages — score changes propagate in 1–2 seconds.",
            "Marquee refresher subscribes globally so even pages without a per-page subscriber stay current.",
        ],
    )

    # ---------- Scorer ----------
    add_heading(doc, "Scorer features", level=1)

    add_section_block(doc,
        "Scorer dashboard (/scorer)",
        "Tournament-day workhorse: filterable match list, tap to score.",
        [
            "Filter chips: Court (1–6), Category (MB / MI / W), Status (Open / Live / Done).",
            "Match table sortable by time, court, category, stage, with team names and live status pill.",
            "Any scorer can score any match (no per-court assignment); each score event records actor identity for audit.",
        ],
    )

    add_section_block(doc,
        "Score entry (/scorer/match/[id])",
        "Optimistic, finger-friendly score input.",
        [
            "Explicit Start match button — no more accidental auto-flip when typing the first score.",
            "Per-game score input: tap +/− for instant updates, or type a value (commits on blur or Enter; Escape reverts).",
            "Optimistic UI — local update is instant; realtime sync continues to work for the same match in another tab.",
            "End game per game (locks that game, advances to the next).",
            "Set winner (Team A / Team B) buttons after games are scored.",
            "Walkover form — pick winner + reason; logs walkover separately in the audit log.",
            "Reset match — clear all games back to 0–0 pending and status back to scheduled.",
            "Locked-match guard: scorers see \"Match is locked, ask an admin\"; admins bypass automatically.",
            "KO matches with unresolved feeders show a friendly \"waiting on prior matches\" banner instead of unusable inputs.",
        ],
    )

    # ---------- Admin ----------
    add_heading(doc, "Admin features", level=1)

    add_section_block(doc,
        "Admin overview (/admin)",
        "Single-glance health of the tournament.",
        [
            "Match counts: total / scheduled / live / completed / walkover / locked.",
            "Group qualifiers confirmed N/M.",
            "Quick links into per-category review and a global \"Re-resolve all brackets\" button.",
        ],
    )

    add_section_block(doc,
        "Per-category review (/admin/categories/[code])",
        "Where the admin transitions group stage → knockout.",
        [
            "Per-group standings with current ranking, tie-break columns, and TOSS badge if a manual tie-break is required.",
            "Qualifier picker: defaults to live ranking's top-2 but admin can override (e.g., for toss-decided ties).",
            "Confirm qualifiers (locks the group's top-2). Unlock to reopen.",
            "Lock all group matches — bulk lock to freeze scoring.",
            "Resolve bracket button — walks every KO match and fills team slots whose feeders are now resolvable. Idempotent and runs automatically when a winner is set.",
        ],
    )

    add_section_block(doc,
        "Match override (/admin/match/[id])",
        "Surgical correction tooling for individual matches.",
        [
            "Same score-entry UI as the scorer view, but force-edit is enabled — admin can change scores and winners even on completed or walkover matches.",
            "Lock / unlock toggle (a lock prevents scorer edits but admin always bypasses).",
            "Full audit log table: timestamp, action, actor role, score deltas, notes, sorted newest first.",
            "Cross-links to the public match view and the scorer view.",
        ],
    )

    add_section_block(doc,
        "User + season management",
        "Who can sign in, and which season is active.",
        [
            "Multi-role: a single email can hold both admin and scorer simultaneously (checkbox-style picker).",
            "Add user, change roles, remove user; status column shows whether they've ever logged in.",
            "Profile auto-syncs role changes so users don't need to re-login.",
            "Create new season (year + name) — appears in the table.",
            "Set active season — only one active at a time, enforced at the DB level.",
            "Per-season branding override available via JSONB column for future season-specific themes.",
        ],
    )

    # ---------- Auth ----------
    add_heading(doc, "Authentication", level=1)
    add_section_block(doc,
        "Email-only direct login",
        "Friction-minimised for an internal corporate event.",
        [
            "Type email, click Sign in — server verifies whitelist, mints a session via the Supabase admin API, sets the cookie. No email is sent at any step.",
            "Whitelist enforced server-side: emails not in allowed_users are rejected before any session machinery runs.",
            "Multi-role on a single user (admin + scorer) supported throughout the stack.",
            "Sign-out button in the header (desktop + mobile drawer).",
            "Session refresh handled by middleware on every request.",
        ],
    )

    # ---------- UI / UX ----------
    add_heading(doc, "UI / UX system", level=1)
    add_section_block(doc,
        "Visual identity",
        None,
        [
            "Navy primary (#1E2A6E) matched to the SBL logo, with warm coral accent (#C0623F).",
            "Cream / warm-ivory background (#FAF9F6) with subtle radial gradients.",
            "Display serif headings (\"Where coworkers become rivals.\") paired with Geist sans body type.",
            "Full dark-mode palette with lifted navy (#6E80C4) for contrast.",
            "Coral-tinted focus rings, smooth 120ms transitions on all interactive elements.",
            "SBL logo wired into nav (desktop + mobile drawer) and centered in the home hero at h-20 / h-28.",
        ],
    )

    add_section_block(doc,
        "Status semantics",
        "Color-coded so live vs done vs scheduled is unmistakable.",
        [
            "Live (in_progress) — coral wash + 4px coral left accent + pulsing LIVE pill.",
            "Completed — green wash + green left accent + FT pill.",
            "Walkover — amber wash + amber left accent + W/O pill.",
            "Scheduled — neutral surface + TBP pill.",
            "Cancelled — dimmed grey.",
            "Applied uniformly to MatchCard (everywhere matches are listed), BracketView, and PredictionCard.",
            "Resolved predictions amplify: bright green wash + glow if your pick was correct, red strike-through if wrong.",
        ],
    )

    add_section_block(doc,
        "Navigation",
        None,
        [
            "Top sticky nav: SBL logo, search button (Cmd-K), category links, courts, brackets, predict, leaderboard, scorer/admin links by role, user menu.",
            "Universal search modal with type-ahead + arrow-key navigation; pre-loaded team + player index.",
            "Mobile sidebar drawer (< 1024px): full-height, scroll-locked, hierarchical sections (Tournament / Brackets / Courts / Play-along / Scorer / Admin) with role-gated visibility.",
            "Breadcrumbs on every detail page (categories, groups, teams, players, matches, courts, bracket, scorer, admin).",
            "Footer with app tagline + active season.",
        ],
    )

    # ---------- Data model ----------
    add_heading(doc, "Data model", level=1)
    add_table(doc,
        ["Table / view", "Purpose"],
        [
            ["seasons", "One row per year (SBL 2026, 2027, …); tracks status + active flag + branding JSONB."],
            ["categories", "MB / MI / W per season with format rules (group_format, ko_format) as JSONB."],
            ["groups", "Group A/B/C/D per category. Includes qualifier_1_team_id, qualifier_2_team_id, qualifiers_locked for admin's manual top-2 confirmation."],
            ["players", "Cross-season identity; deduped on (full_name, company)."],
            ["teams", "Per-season; FK to category + group."],
            ["team_players", "Composite link team ↔ player."],
            ["matches", "Group + KO matches. KO uses team_a_source / team_b_source JSONB feeders that the resolver consumes."],
            ["games", "One row per game in a match (1 for group, 3 for KO)."],
            ["score_events", "Append-only audit log of every score mutation."],
            ["allowed_users", "Whitelist of emails with a roles user_role[] column for multi-role."],
            ["profiles", "Auto-populated from auth.users via trigger; mirrors roles[]."],
            ["participant_profiles", "Display name keyed by localStorage device UUID (no auth)."],
            ["cheers", "Anonymous tap log per match (clap / fire). Powers realtime cheer counters."],
            ["predictions", "One pick per device per KO match; locks when match starts."],
            ["standings_view", "Per-team stats per group: P, W, L, points, sets won/lost, set diff, point diff."],
        ],
    )

    # ---------- Tooling ----------
    add_heading(doc, "Tooling", level=1)
    add_section_block(doc,
        "Scripts",
        None,
        [
            "npm run seed:2026 — imports SBL 2026 fixtures from data/SBL_2026_Fixtures.xlsx into Supabase. --force flag wipes and reseeds.",
            "npm run demo:simulate — populates a realistic mid-tournament state (most groups completed, some matches live, qualifiers confirmed, partial KO progress) for UI/UX testing.",
            "scripts/generate-feature-doc.py — regenerates this document.",
        ],
    )

    add_section_block(doc,
        "Testing + documentation",
        None,
        [
            "TESTING.md — persona-based manual test plan with Participant (10 walkthroughs), Scorer (12), Admin (11), plus a 5-minute end-to-end smoke test recipe and 8 common edge cases.",
            "Repo memory files document architectural decisions and known gotchas (PostgREST view embeds, FK disambiguation, etc.).",
        ],
    )

    # ---------- Out of scope (transparency) ----------
    add_heading(doc, "Deliberate non-features (v1)", level=1)
    add_para(doc,
        "These were considered and explicitly deferred to keep v1 focused.",
        italic=True, color=GREY,
    )
    for b in [
        "Comments / chat on matches — moderation overhead too high for the timeframe.",
        "Photo uploads — storage + content moderation overhead.",
        "Web push notifications — most users are on-site, low ROI for a single-day event.",
        "Mid-event match reschedules from the UI — flagged verbally, app reflects truth via score entry.",
        "Company-vs-company leaderboard — stacking the data was easy, but the UX wasn't worth the ship for v1.",
        "Sound effects on score change — likely annoying.",
        "MVP voting per match — would need anti-spam plumbing.",
    ]:
        add_bullet(doc, b)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
